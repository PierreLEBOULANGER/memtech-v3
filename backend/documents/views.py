"""
Ce fichier définit les vues pour l'API des documents.
Il gère toutes les opérations CRUD sur les documents.
"""

from rest_framework import viewsets, permissions, status
from rest_framework.decorators import authentication_classes
from django.views.decorators.csrf import csrf_exempt
from rest_framework.response import Response
from rest_framework.decorators import action, api_view, permission_classes
from django.shortcuts import get_object_or_404
from .models import Document
from .serializers import DocumentSerializer
from django.core.exceptions import PermissionDenied
from django.conf import settings
import os
from docx import Document as DocxDocument
from projects.models import Project
import jwt  # Librairie PyJWT pour signer le payload OnlyOffice
from .onlyoffice_utils import sign_onlyoffice_payload  # Utilitaire pour signer le payload OnlyOffice
import hashlib
from rest_framework.permissions import AllowAny
from urllib.parse import urlparse
from django.urls import reverse
import base64
import requests
from django.core.files.base import ContentFile

class DocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet pour gérer les opérations CRUD sur les documents.
    Inclut des endpoints personnalisés pour des actions spécifiques.
    """
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        """
        Surcharge la création pour ajouter l'utilisateur actuel comme uploader.
        """
        serializer.save(uploaded_by=self.request.user)

    def get_queryset(self):
        """
        Filtre les documents en fonction du projet si spécifié.
        """
        queryset = Document.objects.all()
        project_id = self.request.query_params.get('project', None)
        if project_id:
            queryset = queryset.filter(project_id=project_id)
        return queryset

    @action(detail=True, methods=['post'])
    def change_status(self, request, pk=None):
        """
        Endpoint personnalisé pour changer le statut d'un document.
        """
        document = self.get_object()
        new_status = request.data.get('status')
        
        if not new_status:
            return Response(
                {'error': 'Le statut est requis'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        document.status = new_status
        document.save()
        
        serializer = self.get_serializer(document)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def change_category(self, request, pk=None):
        """
        Endpoint personnalisé pour changer la catégorie d'un document.
        """
        document = self.get_object()
        new_category = request.data.get('category')
        
        if not new_category:
            return Response(
                {'error': 'La catégorie est requise'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        document.category = new_category
        document.save()
        
        serializer = self.get_serializer(document)
        return Response(serializer.data)

    # @action(detail=True, methods=['get'], url_path='word_url', permission_classes=[AllowAny])
    # def get_or_create_word_file(self, request, pk=None):
    #     """
    #     Endpoint pour obtenir l'URL du fichier Word associé à ce document.
    #     Cette version génère automatiquement une URL HTTPS sur le port 8443 pour OnlyOffice,
    #     et utilise la vue serve_docx pour supprimer le header X-Frame-Options.
    #     """
    #     document = self.get_object()
        
    #     # Si le fichier existe déjà
    #     if document.file and document.file.name.endswith('.docx') and os.path.isfile(document.file.path):
    #         with open(document.file.path, 'rb') as file:
    #             file_content = file.read()
    #             base64_content = base64.b64encode(file_content).decode('utf-8')
                
    #             onlyoffice_payload = {
    #                 "document": {
    #                     "fileType": "docx",
    #                     "key": f"memoire_{document.project.id}_{document.id}",
    #                     "title": "Mémoire technique",
    #                     "url": f"https://courant.eu.ngrok.io/media/memoires/memoire_{document.project.id}_{document.id}.docx"
    #                 },
    #                 "permissions": {
    #                     "edit": True,
    #                     "download": True,
    #                     "print": True
    #                 },
    #                 "editorConfig": {
    #                     "mode": "edit",
    #                     "lang": "fr",
    #                     "callbackUrl": f"https://courant.eu.ngrok.io/api/documents/onlyoffice_callback/",
    #                     "user": {
    #                         "id": str(request.user.id) if request.user.is_authenticated else "1",
    #                         "name": request.user.get_full_name() if request.user.is_authenticated else "Utilisateur"
    #                     }
    #                 }
    #             }
    #             token = sign_onlyoffice_payload(onlyoffice_payload)
    #             return Response({'token': token})

    #     # Si le fichier n'existe pas, on le crée
    #     if DocxDocument is None:
    #         return Response({'error': 'python-docx n\'est pas installé sur le serveur.'}, status=500)

    #     # Créer le document Word vierge
    #     memoires_dir = os.path.join(settings.MEDIA_ROOT, 'memoires')
    #     os.makedirs(memoires_dir, exist_ok=True)
    #     filename = f"memoire_{document.project.id}_{document.id}.docx"
    #     file_path = os.path.join(memoires_dir, filename)

    #     docx = DocxDocument()
    #     docx.add_paragraph(f"Mémoire technique - {document.project.name}")
    #     docx.save(file_path)

    #     # Lire le fichier créé et le convertir en base64
    #     with open(file_path, 'rb') as file:
    #         file_content = file.read()
    #         base64_content = base64.b64encode(file_content).decode('utf-8')

    #     # Associer le fichier au document et sauvegarder
    #     relative_path = os.path.join('memoires', filename)
    #     document.file.name = relative_path
    #     document.save()

    #     # Construction du payload OnlyOffice avec le contenu en base64
    #     onlyoffice_payload = {
    #         "document": {
    #             "fileType": "docx",
    #             "key": f"memoire_{document.project.id}_{document.id}",
    #             "title": "Mémoire technique",
    #             "url": f"https://courant.eu.ngrok.io/media/memoires/memoire_{document.project.id}_{document.id}.docx"
    #         },
    #         "permissions": {
    #             "edit": True,
    #             "download": True,
    #             "print": True
    #         },
    #         "editorConfig": {
    #             "mode": "edit",
    #             "lang": "fr",
    #             "callbackUrl": f"https://courant.eu.ngrok.io/api/documents/onlyoffice_callback/",
    #             "user": {
    #                 "id": str(request.user.id) if request.user.is_authenticated else "1",
    #                 "name": request.user.get_full_name() if request.user.is_authenticated else "Utilisateur"
    #             }
    #         }
    #     }
    #     token = sign_onlyoffice_payload(onlyoffice_payload)
    #     return Response({
    #         'onlyoffice_token': token,
    #         'document_data': base64_content
    #     })

    @action(detail=False, methods=['post'], url_path='create_memoire_technique')
    def create_memoire_technique(self, request):
        """
        Endpoint pour créer un document mémoire technique vierge pour un projet donné.
        Si un document technique existe déjà pour ce projet, le retourne.
        Sinon, crée un document vierge (catégorie 'technical', statut 'draft').
        """

        project_id = request.data.get('project_id')
        if not project_id:
            return Response({'error': 'project_id requis'}, status=400)
        try:
            project = Project.objects.get(id=project_id)
        except Project.DoesNotExist:
            return Response({'error': 'Projet introuvable'}, status=404)
        
        # Vérifier s'il existe déjà un document technique pour ce projet
        existing_doc = Document.objects.filter(project=project, category='technical').first()
        if existing_doc:
            serializer = self.get_serializer(existing_doc)
            return Response({
                'document': serializer.data,
                'document_id': existing_doc.id,
                'message': 'Document technique déjà existant'
            })
    
        
        # Créer un document vierge
        doc = Document.objects.create(
            project=project,
            category='technical',
            status='draft',
            name=f"Mémoire technique - {project.name}",
            description="Mémoire technique du projet",
            uploader=request.user
        )

    
        # Créer le fichier Word vierge
        if DocxDocument is None:
            return Response({'error': 'python-docx n\'est pas installé sur le serveur.'}, status=500)

        # Définir le chemin de stockage
        memoires_dir = os.path.join(settings.MEDIA_ROOT, 'memoires')
        os.makedirs(memoires_dir, exist_ok=True)
        filename = f"memoire_{project.id}_{doc.id}.docx"
        file_path = os.path.join(memoires_dir, filename)

        # Créer un document Word vierge
        docx = DocxDocument()
        docx.add_paragraph(f"Mémoire technique - {project.name}")
        docx.save(file_path)

        # Associer le fichier au document et sauvegarder
        relative_path = os.path.join('memoires', filename)
        doc.file.name = relative_path
        doc.save()

        serializer = self.get_serializer(doc)
        return Response({
            'document': serializer.data,
            'document_id': doc.id,
            'message': 'Document technique créé avec succès'
        })


@csrf_exempt
@api_view(['POST'])
@authentication_classes([])  
@permission_classes([AllowAny])
def onlyoffice_callback(request):
    print("Callback OnlyOffice reçu :", request.data)  # Pour debug
    print("Headers reçus :", request.headers)  # Pour voir les headers
    print("Méthode :", request.method)  # Pour voir la méthode HTTP

    data = request.data
    status = data.get('status')
    file_url = data.get('url')
    key = data.get('key')

    print(f"Status: {status}, URL: {file_url}, Key: {key}")  # Pour voir les détails

    # Statuts OnlyOffice : 2 = sauvegarde, 6 = sauvegarde forcée
    if status in [2, 6] and file_url and key:
        # Extraire l'id du document à partir de la clé (ex: "memoire_77_20")
        try:
            # Ici, on suppose que la clé est de la forme "memoire_{project_id}_{doc_id}"
            doc_id = int(key.split('_')[-1])
            from .models import Document
            document = Document.objects.get(id=doc_id)
        except Exception as e:
            print("Erreur récupération document:", e)
            return Response({'error': 'Document introuvable'}, status=404)

        # Télécharger le fichier modifié depuis OnlyOffice
        try:
            r = requests.get(file_url)
            r.raise_for_status()
            # Remplacer le fichier existant
            filename = document.file.name.split('/')[-1]
            document.file.save(filename, ContentFile(r.content), save=True)
            print("Fichier sauvegardé !")
        except Exception as e:
            print("Erreur lors du téléchargement/sauvegarde :", e)
            return Response({'error': 'Erreur lors de la sauvegarde'}, status=500)

    return Response({"error":0})

# @api_view(['GET'])
# def word_url(request, pk):
#     """
#     Vue API qui retourne l'URL du document Word pour OnlyOffice et le token JWT.
#     Utilité : Permettre au front de récupérer l'URL et le token à transmettre à OnlyOffice.
#     """
#     document = get_object_or_404(Document, pk=pk)
#     url = f"/api/documents/media/{document.file.name}"
#     # Génération du token JWT OnlyOffice (si besoin)
#     onlyoffice_payload = {
#         "document": {
#             "fileType": "docx",
#             "key": hashlib.sha1(url.encode('utf-8')).hexdigest(),
#             "title": "Mémoire technique",
#             "url": url,
#         },
#         "documentType": "word",
#         "editorConfig": {
#             "mode": "edit",
#             "lang": "fr",
#             "callbackUrl": "https://courant.eu.ngrok.io/api/documents/onlyoffice_callback/",
#             "user": {
#                 "id": str(request.user.id) if request.user.is_authenticated else "1",
#                 "name": request.user.get_full_name() if request.user.is_authenticated else "Utilisateur"
#             }
#         },
#         "permissions": {
#             "edit": True,
#             "download": True,
#             "print": True,
#             "review": True,
#         },
#         "height": "100%",
#         "width": "100%",
#     }
#     token = sign_onlyoffice_payload(onlyoffice_payload)
#     return Response({"url": url, "onlyoffice_token": token}, status=status.HTTP_200_OK)
